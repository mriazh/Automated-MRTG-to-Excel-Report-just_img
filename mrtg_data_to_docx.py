import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ========== KONFIGURASI ==========
FOLDER_INPUT = "MRTG-Data"          # Folder tempat subfolder tanggal
OUTPUT_DOCX = "MRTG_Report.docx"    # Nama file output
LIST_FILE = "list_mrtg_data_to_docx.txt"   # File daftar urutan (bisa juga hardcode di bawah)

# Ukuran gambar di Word (lebar dalam inci)
IMAGE_WIDTH = Inches(1.5)
# Jumlah kolom dalam tabel (agar muat 20 gambar per halaman)
COLUMNS = 5

# ========== BACA DAFTAR DARI FILE ==========
def baca_daftar(filepath):
    """
    Baca file list_mrtg_data_to_docx.txt, return list of tuples:
    (nomor, tipe, id)
    tipe: 'SID' atau 'Graph-title'
    id: string (misal '4700001-0021497479' atau '3598')
    """
    items = []
    with open(filepath, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            # Format: "1. SID : 4700001-0021497479"
            parts = line.split('.', 1)
            if len(parts) != 2:
                continue
            nomor = parts[0].strip()
            rest = parts[1].strip()
            if rest.startswith('SID : '):
                tipe = 'SID'
                id_val = rest.replace('SID : ', '').strip()
            elif rest.startswith('Graph-title : '):
                tipe = 'Graph-title'
                id_val = rest.replace('Graph-title : ', '').strip()
            else:
                continue
            items.append((nomor, tipe, id_val))
    return items

# ========== CARI FILE GAMBAR DI FOLDER TANGGAL ==========
def cari_gambar(folder_tanggal, tipe, id_val, tanggal_str):
    """
    Cari file gambar sesuai tipe dan id.
    - Untuk SID: cari file dengan nama "MRTG_<id>.png" (case sensitive)
    - Untuk Graph-title: cari file dengan nama "MRTG_<id>_<tanggal_str>.png"
    Return path file jika ditemukan, None jika tidak.
    """
    if tipe == 'SID':
        nama_target = f"MRTG_{id_val}.png"
        path = os.path.join(folder_tanggal, nama_target)
        if os.path.exists(path):
            return path
        # fallback: coba case-insensitive (jika ada)
        for f in os.listdir(folder_tanggal):
            if f.lower() == nama_target.lower():
                return os.path.join(folder_tanggal, f)
        return None
    else:  # Graph-title
        nama_target = f"MRTG_{id_val}_{tanggal_str}.png"
        path = os.path.join(folder_tanggal, nama_target)
        if os.path.exists(path):
            return path
        # fallback: cari file yang mengandung id dan tanggal
        for f in os.listdir(folder_tanggal):
            if f.startswith(f"MRTG_{id_val}") and tanggal_str in f:
                return os.path.join(folder_tanggal, f)
        return None

# ========== TAMBAHKAN GAMBAR KE DALAM SEL TABEL ==========
def tambah_gambar_ke_sel(sel, image_path, nomor, label):
    """
    Masukkan gambar ke dalam sel tabel, lalu tambahkan caption.
    """
    try:
        # Cek ukuran gambar asli (optional)
        with Image.open(image_path) as img:
            # hitung proporsi agar tidak terlalu besar
            pass
        # Tambahkan gambar dengan lebar tetap
        run = sel.paragraphs[0].add_run()
        run.add_picture(image_path, width=IMAGE_WIDTH)
        # Atur alignment gambar ke tengah
        sel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Tambahkan baris baru untuk caption
        p = sel.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = f"{nomor}. {label}"
        p.add_run(caption).font.size = Inches(0.1)
    except Exception as e:
        sel.paragraphs[0].text = f"Error: {os.path.basename(image_path)}"
        print(f"  Gagal memuat {image_path}: {e}")

# ========== PROSES SEMUA TANGGAL ==========
def proses_semua_tanggal(doc, items, root_folder):
    # Dapatkan semua subfolder yang namanya angka (YYYYMMDD)
    subfolders = [f for f in os.listdir(root_folder) 
                  if os.path.isdir(os.path.join(root_folder, f)) and f.isdigit()]
    subfolders.sort()  # urutkan dari tanggal awal ke akhir

    for tanggal_folder in subfolders:
        full_path = os.path.join(root_folder, tanggal_folder)
        print(f"Memproses tanggal: {tanggal_folder}")

        # Tambahkan heading tanggal
        doc.add_heading(f"Tanggal: {tanggal_folder[:4]}-{tanggal_folder[4:6]}-{tanggal_folder[6:]}", level=1)

        # Buat tabel dengan jumlah baris = ceil(20/COLUMNS), kolom = COLUMNS
        rows = (len(items) + COLUMNS - 1) // COLUMNS
        table = doc.add_table(rows=rows, cols=COLUMNS)
        table.style = 'Table Grid'  # biar ada border

        # Loop setiap item, isi ke sel tabel secara berurutan (row-major)
        for idx, (nomor, tipe, id_val) in enumerate(items):
            row = idx // COLUMNS
            col = idx % COLUMNS
            sel = table.cell(row, col)

            # Cari gambar
            label = f"{tipe} {id_val}"
            if tipe == 'SID':
                image_path = cari_gambar(full_path, tipe, id_val, tanggal_folder)
            else:
                image_path = cari_gambar(full_path, tipe, id_val, tanggal_folder)

            if image_path:
                tambah_gambar_ke_sel(sel, image_path, nomor, label)
            else:
                sel.paragraphs[0].text = f"{nomor}. {label}\n(Tidak ditemukan)"
                print(f"  WARNING: Gambar tidak ditemukan untuk {label} di {tanggal_folder}")

        # Beri jarak antar tanggal
        doc.add_page_break()

# ========== MAIN ==========
def main():
    print("=" * 60)
    print("MRTG to DOCX - Generate Report per Tanggal")
    print("=" * 60)

    # Baca daftar urutan
    if not os.path.exists(LIST_FILE):
        print(f"File daftar '{LIST_FILE}' tidak ditemukan!")
        return
    items = baca_daftar(LIST_FILE)
    print(f"Daftar berisi {len(items)} item (1..20)")

    # Pastikan folder input ada
    if not os.path.exists(FOLDER_INPUT):
        print(f"Folder input '{FOLDER_INPUT}' tidak ditemukan!")
        return

    # Buat dokumen Word
    doc = Document()
    doc.add_heading('Laporan Grafik MRTG', level=0)

    # Proses semua tanggal
    proses_semua_tanggal(doc, items, FOLDER_INPUT)

    # Simpan
    doc.save(OUTPUT_DOCX)
    print(f"\n✅ Selesai! File DOCX disimpan sebagai: {OUTPUT_DOCX}")
    print(f"   Buka dengan Microsoft Word dan lihat hasilnya.")

if __name__ == "__main__":
    main()